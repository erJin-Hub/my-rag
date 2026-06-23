from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return file_bytes.decode("utf-8-sig")
    if suffix == ".pdf":
        return extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return extract_docx_text(file_bytes)
    raise ValueError(f"不支持的文件类型: {suffix}，只支持 {sorted(SUPPORTED_SUFFIXES)}")


def extract_text_from_path(file_path: Path) -> str:
    return extract_text_from_bytes(file_path.read_bytes(), file_path.name)


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(f"[第 {page_number} 页]\n{text}")
    return "\n\n".join(pages)


def extract_docx_text(file_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))

    return "\n".join(paragraphs + table_rows)
